"""Python wrapper to execute report_builder.js or fall back to python-docx."""

import json
import logging
import os
import shutil
import subprocess
import tempfile

logger = logging.getLogger("gap_analyzer")

MODEL = os.getenv("CLAUDE_MODEL", "claude-opus-4-6")


def build_report(cache_data: dict, output_path: str) -> bool:
    """Build Word document report. Uses Node.js if available, else python-docx fallback."""
    node_path = shutil.which("node")
    script_dir = os.path.dirname(os.path.abspath(__file__))

    if node_path:
        return _build_with_node(cache_data, output_path, node_path, script_dir)
    else:
        logger.info("[DECISION] Node.js not found — falling back to python-docx report builder")
        return _build_with_python_docx(cache_data, output_path)


def _build_with_node(cache_data: dict, output_path: str, node_path: str, script_dir: str) -> bool:
    """Build report using Node.js docx package."""
    # Write cache to temp JSON file for Node script to read
    tmp_json = tempfile.NamedTemporaryFile(
        mode="w", suffix=".json", delete=False, encoding="utf-8"
    )
    try:
        json.dump(cache_data, tmp_json, ensure_ascii=False, indent=2)
        tmp_json.close()

        js_script = os.path.join(script_dir, "report_builder.js")

        # Ensure npm packages are installed
        pkg_dir = os.path.dirname(script_dir)  # gap_analyzer root
        pkg_json = os.path.join(pkg_dir, "package.json")
        node_modules = os.path.join(pkg_dir, "node_modules")
        if os.path.exists(pkg_json) and not os.path.exists(node_modules):
            npm_path = shutil.which("npm")
            if npm_path:
                logger.info("[INFO] Installing Node.js dependencies...")
                subprocess.run(
                    [npm_path, "install"],
                    cwd=pkg_dir,
                    capture_output=True,
                    timeout=60,
                )
            else:
                logger.warning("[WARN] npm not found — skipping Node.js dependency install")
                return _build_with_python_docx(cache_data, output_path)

        try:
            result = subprocess.run(
                [node_path, js_script, tmp_json.name, output_path],
                capture_output=True,
                text=True,
                timeout=300,
                cwd=pkg_dir,
            )
        except subprocess.TimeoutExpired:
            logger.error("[ERROR] Node.js report builder timed out after 5 minutes")
            return _build_with_python_docx(cache_data, output_path)

        if result.returncode == 0:
            logger.info(f"[INFO] Word document generated: {output_path}")
            return True
        else:
            logger.error(f"[ERROR] Node.js report builder failed: {result.stderr}")
            logger.info("[DECISION] Falling back to python-docx report builder")
            return _build_with_python_docx(cache_data, output_path)

    finally:
        try:
            os.unlink(tmp_json.name)
        except OSError:
            pass


def _build_with_python_docx(cache_data: dict, output_path: str) -> bool:
    """Fallback: build report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        navy = RGBColor(0x1B, 0x2A, 0x4A)
        slate = RGBColor(0x2C, 0x3E, 0x50)
        blue = RGBColor(0x2E, 0x86, 0xC1)

        metadata = cache_data.get("metadata", {})
        project = cache_data.get("project_profile", {})
        companies = cache_data.get("companies", {})
        synthesis = cache_data.get("synthesis", {})

        total = len(companies)

        # ─── COVER PAGE ───
        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_heading("Competitive Intelligence Report", level=0)
        for run in title.runs:
            run.font.color.rgb = navy
            run.font.size = Pt(28)
        subtitle = doc.add_paragraph(
            f"Ecommerce Ecosystem Analysis — {total} Companies"
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Date: {metadata.get('last_updated', 'N/A')}")
        doc.add_paragraph("Version: 1.0 — Confidential")
        doc.add_page_break()

        # ─── TABLE OF CONTENTS placeholder ───
        h = doc.add_heading("Table of Contents", level=1)
        for run in h.runs:
            run.font.color.rgb = navy
        doc.add_paragraph(
            "(Update this field in Word: right-click → Update Field)",
            style="Normal"
        )
        doc.add_page_break()

        # ─── SECTION 1: EXECUTIVE SUMMARY ───
        h = doc.add_heading("1. Executive Summary", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        # Project profile
        h2 = doc.add_heading("Project Profile", level=2)
        for run in h2.runs:
            run.font.color.rgb = slate
        services = project.get("services_and_features", [])
        if services:
            doc.add_paragraph(f"Services: {', '.join(services[:10])}")
        doc.add_paragraph(f"Business Model: {project.get('business_model', 'N/A')}")
        doc.add_paragraph(f"Target Audience: {', '.join(project.get('target_audience', ['N/A']))}")

        # Executive summary from synthesis
        exec_sum = synthesis.get("executive_summary", "")
        if exec_sum:
            doc.add_paragraph(exec_sum)

        # Quick wins
        quick_wins = synthesis.get("quick_wins", [])
        if quick_wins:
            h2 = doc.add_heading("Quick Wins", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            hdr[0].text = "Action"
            hdr[1].text = "Impact"
            for qw in quick_wins:
                row = table.add_row().cells
                row[0].text = qw
                row[1].text = "High"

        # Strategic priorities
        priorities = synthesis.get("strategic_priorities", [])
        if priorities:
            h2 = doc.add_heading("Strategic Priorities", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            for i, p in enumerate(priorities, 1):
                doc.add_paragraph(f"{i}. {p}")

        doc.add_page_break()

        # ─── SECTION 2: KEY FINDINGS BY DIMENSION ───
        h = doc.add_heading("2. Key Findings by Dimension", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        missing_features = synthesis.get("top_missing_features", [])
        if missing_features:
            h2 = doc.add_heading("Top Missing Features", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Feature"
            hdr[1].text = "Seen At"
            hdr[2].text = "Priority"
            hdr[3].text = "Recommendation"
            for mf in missing_features:
                if isinstance(mf, dict):
                    row = table.add_row().cells
                    row[0].text = mf.get("feature", "")
                    row[1].text = ", ".join(mf.get("seen_at", []))
                    row[2].text = mf.get("priority", "")
                    row[3].text = mf.get("recommendation", "")

        content_gaps = synthesis.get("content_strategy_gaps", [])
        if content_gaps:
            h2 = doc.add_heading("Content Strategy Gaps", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            for cg in content_gaps:
                doc.add_paragraph(f"  {cg}", style="List Bullet")

        biz_model = synthesis.get("business_model_enhancements", [])
        if biz_model:
            h2 = doc.add_heading("Business Model Enhancements", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            for bm in biz_model:
                doc.add_paragraph(f"  {bm}", style="List Bullet")

        doc.add_page_break()

        # ─── SECTION 3: FINDINGS BY CATEGORY ───
        h = doc.add_heading("3. Findings by Category", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        cat_insights = synthesis.get("category_by_category_insights", [])
        if cat_insights:
            for ci in cat_insights:
                if isinstance(ci, dict):
                    h2 = doc.add_heading(ci.get("category", "Unknown"), level=2)
                    for run in h2.runs:
                        run.font.color.rgb = slate
                    players = ci.get("key_players", [])
                    if players:
                        doc.add_paragraph(f"Key Players: {', '.join(players)}")
                    doc.add_paragraph(f"Patterns: {ci.get('dominant_patterns', 'N/A')}")
                    doc.add_paragraph(f"Project Gaps: {ci.get('project_gaps', 'N/A')}")
                    recs = ci.get("recommendations", [])
                    for r in recs:
                        doc.add_paragraph(f"  {r}", style="List Bullet")

        doc.add_page_break()

        # ─── SECTION 4: COMPANY-BY-COMPANY ANALYSIS ───
        h = doc.add_heading("4. Company-by-Company Analysis", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        # Group by category
        by_category = {}
        for domain, data in companies.items():
            analysis = data.get("analysis", data)
            cat = analysis.get("category", "Uncategorised") or "Uncategorised"
            if cat not in by_category:
                by_category[cat] = []
            by_category[cat].append((domain, analysis))

        for category, entries in sorted(by_category.items()):
            h2 = doc.add_heading(f"{category} ({len(entries)} companies)", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate

            for domain, analysis in entries:
                name = analysis.get("company_name", domain)
                url = analysis.get("url", "")
                verdict = analysis.get("one_line_verdict", "")
                niche = analysis.get("niche", "")

                h3 = doc.add_heading(f"{name}", level=3)
                for run in h3.runs:
                    run.font.color.rgb = blue

                doc.add_paragraph(f"URL: {url}  |  Category: {category}  |  Niche: {niche}")
                if verdict:
                    p = doc.add_paragraph()
                    p.add_run("Verdict: ").bold = True
                    p.add_run(verdict)

                # Dimension table
                dims = [
                    ("Functionality", "dim2_functionality_tech"),
                    ("Content", "dim3_content_messaging"),
                    ("Services", "dim4_services_products"),
                    ("Business Model", "dim5_business_model"),
                ]
                table = doc.add_table(rows=1, cols=3)
                hdr = table.rows[0].cells
                hdr[0].text = "Dimension"
                hdr[1].text = "What They Do"
                hdr[2].text = "Gap For Your Project"
                for dim_label, dim_key in dims:
                    dim_data = analysis.get(dim_key, {})
                    if isinstance(dim_data, dict):
                        summary_parts = []
                        for k, v in dim_data.items():
                            if k != "gap_for_your_project" and isinstance(v, str):
                                summary_parts.append(v)
                        row = table.add_row().cells
                        row[0].text = dim_label
                        row[1].text = "; ".join(summary_parts[:3])[:200]
                        row[2].text = str(dim_data.get("gap_for_your_project", "N/A"))[:200]

                # Opportunities, ideas, watch-outs
                opps = analysis.get("top_opportunities", [])
                if opps:
                    p = doc.add_paragraph()
                    p.add_run("Top Opportunities:").bold = True
                    for o in opps:
                        doc.add_paragraph(f"  {o}", style="List Bullet")

                ideas = analysis.get("value_add_ideas", [])
                if ideas:
                    p = doc.add_paragraph()
                    p.add_run("Value-Add Ideas:").bold = True
                    for i_item in ideas:
                        doc.add_paragraph(f"  {i_item}", style="List Bullet")

                watch = analysis.get("watch_out_for", [])
                if watch:
                    p = doc.add_paragraph()
                    p.add_run("Watch Out For:").bold = True
                    for w in watch:
                        doc.add_paragraph(f"  {w}", style="List Bullet")

        doc.add_page_break()

        # ─── SECTION 5: MASTER OPPORTUNITY MATRIX ───
        h = doc.add_heading("5. Master Opportunity Matrix", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        # Aggregate all opportunities
        all_opps = []
        for domain, data in companies.items():
            analysis = data.get("analysis", data)
            name = analysis.get("company_name", domain)
            for opp in analysis.get("top_opportunities", []):
                all_opps.append({"opportunity": opp, "company": name})

        if all_opps:
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Opportunity"
            hdr[1].text = "Seen At"
            hdr[2].text = "Priority"
            # Show top 50
            for item in all_opps[:50]:
                row = table.add_row().cells
                row[0].text = item["opportunity"][:150]
                row[1].text = item["company"]
                row[2].text = "Medium"

        doc.add_page_break()

        # ─── SECTION 6: STRATEGIC RECOMMENDATIONS ───
        h = doc.add_heading("6. Strategic Recommendations", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        if missing_features:
            h2 = doc.add_heading("Top Missing Features (Detailed)", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            for mf in missing_features:
                if isinstance(mf, dict):
                    doc.add_paragraph(
                        f"{mf.get('feature', '')} — {mf.get('recommendation', '')} "
                        f"(Seen at: {', '.join(mf.get('seen_at', []))})"
                    )

        niches = synthesis.get("underserved_niches", [])
        if niches:
            h2 = doc.add_heading("Underserved Niches to Target", level=2)
            for run in h2.runs:
                run.font.color.rgb = slate
            for n in niches:
                doc.add_paragraph(f"  {n}", style="List Bullet")

        doc.add_page_break()

        # ─── SECTION 7: RISKS ───
        h = doc.add_heading("7. Risks If Not Addressed", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        risks = synthesis.get("risks_if_not_addressed", [])
        if risks:
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Risk"
            hdr[1].text = "Impact"
            hdr[2].text = "Mitigation"
            for risk in risks:
                row = table.add_row().cells
                row[0].text = str(risk)[:200]
                row[1].text = "High"
                row[2].text = "See recommendations"

        doc.add_page_break()

        # ─── APPENDIX A: COMPANY INDEX ───
        h = doc.add_heading("Appendix A: Full Company Index", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Company"
        hdr[1].text = "URL"
        hdr[2].text = "Category"
        hdr[3].text = "Niche"
        hdr[4].text = "Status"
        for domain, data in sorted(companies.items()):
            analysis = data.get("analysis", data)
            status = data.get("status", "success")
            row = table.add_row().cells
            row[0].text = analysis.get("company_name", domain)[:50]
            row[1].text = analysis.get("url", domain)[:60]
            row[2].text = analysis.get("category", "")[:30]
            row[3].text = analysis.get("niche", "")[:30]
            row[4].text = status

        doc.add_page_break()

        # ─── APPENDIX B: DATA SOURCES ───
        h = doc.add_heading("Appendix B: Data Sources", level=1)
        for run in h.runs:
            run.font.color.rgb = navy

        runs_list = metadata.get("runs", [])
        doc.add_paragraph(f"Analysis Date: {metadata.get('last_updated', 'N/A')}")
        doc.add_paragraph(f"Companies File: {metadata.get('companies_file', 'N/A')}")
        doc.add_paragraph(f"Specs File: {metadata.get('specs_file', 'N/A')}")
        doc.add_paragraph(f"Claude Model: {MODEL}")
        doc.add_paragraph(f"Total Companies: {total}")
        if runs_list:
            doc.add_paragraph(f"Total Runs: {len(runs_list)}")
            for r in runs_list:
                doc.add_paragraph(
                    f"  Run {r.get('run_id', 'N/A')}: "
                    f"{r.get('completed', 0)} completed, "
                    f"{r.get('failed', 0)} failed"
                )

        # Save
        os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
        doc.save(output_path)
        logger.info(f"[INFO] Report saved (python-docx): {output_path}")
        return True

    except Exception as e:
        logger.error(f"[ERROR] python-docx report generation failed: {e}")
        # Last resort: plain text fallback
        return _build_plaintext_fallback(cache_data, output_path)


def _build_plaintext_fallback(cache_data: dict, output_path: str) -> bool:
    """Last resort: save as plain text."""
    base, _ = os.path.splitext(output_path)
    txt_path = base + "_fallback.txt"
    try:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("COMPETITIVE INTELLIGENCE REPORT — PLAIN TEXT FALLBACK\n")
            f.write("=" * 60 + "\n\n")

            synthesis = cache_data.get("synthesis", {})
            f.write("EXECUTIVE SUMMARY:\n")
            f.write(synthesis.get("executive_summary", "N/A") + "\n\n")

            companies = cache_data.get("companies", {})
            f.write(f"COMPANIES ANALYSED: {len(companies)}\n\n")

            for domain, data in companies.items():
                analysis = data.get("analysis", data)
                f.write(f"\n--- {analysis.get('company_name', domain)} ---\n")
                f.write(f"URL: {analysis.get('url', '')}\n")
                f.write(f"Verdict: {analysis.get('one_line_verdict', 'N/A')}\n")
                opps = analysis.get("top_opportunities", [])
                if opps:
                    f.write("Opportunities:\n")
                    for o in opps:
                        f.write(f"  - {o}\n")

        logger.info(f"[INFO] Fallback text report saved: {txt_path}")
        return True
    except Exception as e:
        logger.error(f"[ERROR] Even plaintext fallback failed: {e}")
        return False
